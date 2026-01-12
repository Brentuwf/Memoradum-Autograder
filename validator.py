import re
from datetime import datetime
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from enum import Enum

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx library not installed")
    print("Install it using: pip install python-docx")
    exit(1)

class Severity(Enum):
    CRITICAL = "CRITICAL"
    WARNING = "WARNING"
    INFO = "INFO"

@dataclass
class ValidationIssue:
    severity: Severity
    section: str
    message: str
    paragraph_number: Optional[int] = None

class MemorandumValidator:
    def __init__(self):
        self.issues: List[ValidationIssue] = []
        self.document: Optional[Document] = None
        self.paragraphs: List = []
        
    def validate_file(self, file_path: str) -> Tuple[bool, List[ValidationIssue]]:
        """Main validation entry point for .docx files"""
        self.issues = []
        
        if not file_path.endswith('.docx'):
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "File",
                "File must be a .docx document"
            ))
            return False, self.issues
        
        try:
            self.document = Document(file_path)
            self.paragraphs = [p for p in self.document.paragraphs]
            
            # Run all validation checks
            self._validate_date()
            self._validate_memorandum_for()
            self._validate_from_line()
            self._validate_subject_line()
            self._validate_body_paragraphs()
            self._validate_signature_block()
            self._validate_attachments()
            self._validate_formatting()
            
            # Determine overall pass/fail
            critical_issues = [i for i in self.issues if i.severity == Severity.CRITICAL]
            passed = len(critical_issues) == 0
            
            return passed, self.issues
            
        except FileNotFoundError:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "File",
                f"File not found: {file_path}"
            ))
            return False, self.issues
        except Exception as e:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "File",
                f"Error reading file: {str(e)}"
            ))
            return False, self.issues
    
    def _find_paragraph_index(self, pattern: str, start: int = 0) -> int:
        """Find the first paragraph matching a pattern"""
        for i in range(start, len(self.paragraphs)):
            if re.search(pattern, self.paragraphs[i].text, re.IGNORECASE):
                return i
        return -1
    
    def _get_text_by_index(self, index: int) -> str:
        """Safely get paragraph text by index"""
        if 0 <= index < len(self.paragraphs):
            return self.paragraphs[index].text.strip()
        return ""
    
    def _validate_date(self):
        """Validate date format at top of memo (DD Month YYYY)"""
        if not self.paragraphs:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Date",
                "Document is empty"
            ))
            return
        
        # Look for date in first few paragraphs
        date_pattern = r'^\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*$'
        
        found = False
        for i in range(min(5, len(self.paragraphs))):
            text = self.paragraphs[i].text.strip()
            if re.match(date_pattern, text):
                found = True
                # Validate it's a real date
                try:
                    datetime.strptime(text, '%d %B %Y')
                except ValueError:
                    self.issues.append(ValidationIssue(
                        Severity.WARNING,
                        "Date",
                        f"Date format appears correct but may be invalid: {text}",
                        i + 1
                    ))
                break
        
        if not found:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Date",
                "Missing or incorrectly formatted date. Expected format: 'DD Month YYYY' (e.g., '12 March 2025')"
            ))
    
    def _validate_memorandum_for(self):
        """Validate MEMORANDUM FOR line"""
        idx = self._find_paragraph_index(r'MEMORANDUM FOR')
        
        if idx == -1:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Header",
                "Missing 'MEMORANDUM FOR' line"
            ))
        else:
            text = self._get_text_by_index(idx)
            if not text.startswith('MEMORANDUM FOR'):
                self.issues.append(ValidationIssue(
                    Severity.WARNING,
                    "Header",
                    "'MEMORANDUM FOR' should be at the start of the line",
                    idx + 1
                ))
    
    def _validate_from_line(self):
        """Validate FROM: line"""
        idx = self._find_paragraph_index(r'FROM:')
        
        if idx == -1:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Header",
                "Missing 'FROM:' line"
            ))
        else:
            text = self._get_text_by_index(idx)
            if not re.match(r'^FROM:\s+.+', text):
                self.issues.append(ValidationIssue(
                    Severity.WARNING,
                    "Header",
                    "FROM: line should be followed by sender information",
                    idx + 1
                ))
    
    def _validate_subject_line(self):
        """Validate SUBJECT: line"""
        idx = self._find_paragraph_index(r'SUBJECT:')
        
        if idx == -1:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Header",
                "Missing 'SUBJECT:' line"
            ))
        else:
            text = self._get_text_by_index(idx)
            if not re.match(r'^SUBJECT:\s+.+', text):
                self.issues.append(ValidationIssue(
                    Severity.WARNING,
                    "Header",
                    "SUBJECT: line should be followed by subject text",
                    idx + 1
                ))
    
    def _validate_body_paragraphs(self):
        """Validate body paragraph numbering (1., 2., 3., etc.)"""
        # Find where body starts (after SUBJECT line)
        subject_idx = self._find_paragraph_index(r'SUBJECT:')
        if subject_idx == -1:
            return
        
        # Find where signature block starts
        sig_idx = self._find_paragraph_index(r'//SIGNED//', subject_idx)
        if sig_idx == -1:
            # No signature block found, search to end
            sig_idx = len(self.paragraphs)
        
        # Look for numbered paragraphs
        paragraph_pattern = r'^\d+\.\s+'
        found_paragraphs = []
        
        for i in range(subject_idx + 1, sig_idx):
            text = self.paragraphs[i].text.strip()
            match = re.match(r'^(\d+)\.\s+', text)
            if match:
                para_num = int(match.group(1))
                found_paragraphs.append((para_num, i))
        
        if not found_paragraphs:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Body",
                "No numbered paragraphs found. Body paragraphs should be numbered (1., 2., 3., etc.)"
            ))
            return
        
        # Validate sequential numbering
        expected = 1
        for para_num, para_idx in found_paragraphs:
            if para_num != expected:
                self.issues.append(ValidationIssue(
                    Severity.WARNING,
                    "Body",
                    f"Paragraph numbering issue: Expected {expected}, found {para_num}",
                    para_idx + 1
                ))
            expected = para_num + 1
        
        # Info message about paragraph count
        self.issues.append(ValidationIssue(
            Severity.INFO,
            "Body",
            f"Found {len(found_paragraphs)} numbered paragraph(s)"
        ))
    
    def _validate_signature_block(self):
        """Validate signature block format"""
        sig_idx = self._find_paragraph_index(r'//SIGNED//')
        
        if sig_idx == -1:
            self.issues.append(ValidationIssue(
                Severity.CRITICAL,
                "Signature",
                "Missing '//SIGNED//' marker in signature block"
            ))
            return
        
        # Check for name, rank, branch after //SIGNED//
        # Typically 1-3 paragraphs after //SIGNED//
        name_found = False
        for i in range(sig_idx + 1, min(sig_idx + 5, len(self.paragraphs))):
            text = self.paragraphs[i].text.strip()
            # Look for pattern: Name, Rank, Branch (e.g., "Snuff A. Brown, Colonel, USAF")
            if re.search(r'.+,\s*(?:Colonel|Lt Col|Major|Captain|Lieutenant|General|Brig Gen|Maj Gen|Lt Gen|Col|Capt|Lt|Gen),\s*(?:USAF|USSF)', text):
                name_found = True
                break
        
        if not name_found:
            self.issues.append(ValidationIssue(
                Severity.WARNING,
                "Signature",
                "Signature block may be incomplete. Expected format: 'Name, Rank, Branch'",
                sig_idx + 1
            ))
    
    def _validate_attachments(self):
        """Validate attachments section if present"""
        # Look for "Attachments:" or "Attachment:"
        attach_idx = self._find_paragraph_index(r'Attachments?:')
        
        if attach_idx == -1:
            # No attachments section - this is valid
            self.issues.append(ValidationIssue(
                Severity.INFO,
                "Attachments",
                "No attachments section found (optional)"
            ))
            return
        
        # If attachments section exists, validate tabs are listed
        tab_pattern = r'Tab\s+\d+'
        found_tabs = []
        
        for i in range(attach_idx + 1, len(self.paragraphs)):
            text = self.paragraphs[i].text.strip()
            if re.match(tab_pattern, text):
                found_tabs.append(text)
            elif text and not re.match(tab_pattern, text):
                # Non-empty line that's not a tab - end of attachment section
                break
        
        if not found_tabs:
            self.issues.append(ValidationIssue(
                Severity.WARNING,
                "Attachments",
                "Attachments section found but no tabs listed",
                attach_idx + 1
            ))
        else:
            self.issues.append(ValidationIssue(
                Severity.INFO,
                "Attachments",
                f"Found {len(found_tabs)} attachment tab(s)"
            ))
    
    def _validate_formatting(self):
        """Validate basic formatting requirements"""
        if not self.paragraphs:
            return
        
        # Check for common formatting issues
        formatting_issues = []
        
        # Check margins (standard is 1 inch)
        sections = self.document.sections
        if sections:
            section = sections[0]
            # Convert to inches
            top_margin = section.top_margin.inches if section.top_margin else 0
            bottom_margin = section.bottom_margin.inches if section.bottom_margin else 0
            left_margin = section.left_margin.inches if section.left_margin else 0
            right_margin = section.right_margin.inches if section.right_margin else 0
            
            # Allow small tolerance (0.1 inch)
            if abs(top_margin - 1.0) > 0.1:
                self.issues.append(ValidationIssue(
                    Severity.INFO,
                    "Formatting",
                    f"Top margin is {top_margin:.2f} inches (standard is 1.0 inch)"
                ))
            
            if abs(left_margin - 1.0) > 0.1:
                self.issues.append(ValidationIssue(
                    Severity.INFO,
                    "Formatting",
                    f"Left margin is {left_margin:.2f} inches (standard is 1.0 inch)"
                ))


def generate_report(passed: bool, issues: List[ValidationIssue]) -> str:
    """Generate a human-readable validation report"""
    report = []
    report.append("=" * 70)
    report.append("DAF MEMORANDUM VALIDATION REPORT")
    report.append("=" * 70)
    report.append("")
    
    if passed:
        report.append("✓ VALIDATION PASSED")
    else:
        report.append("✗ VALIDATION FAILED")
    
    report.append("")
    
    # Group issues by severity
    critical = [i for i in issues if i.severity == Severity.CRITICAL]
    warnings = [i for i in issues if i.severity == Severity.WARNING]
    info = [i for i in issues if i.severity == Severity.INFO]
    
    if critical:
        report.append("CRITICAL ISSUES:")
        report.append("-" * 70)
        for issue in critical:
            para_info = f" (Paragraph {issue.paragraph_number})" if issue.paragraph_number else ""
            report.append(f"  [{issue.section}]{para_info}: {issue.message}")
        report.append("")
    
    if warnings:
        report.append("WARNINGS:")
        report.append("-" * 70)
        for issue in warnings:
            para_info = f" (Paragraph {issue.paragraph_number})" if issue.paragraph_number else ""
            report.append(f"  [{issue.section}]{para_info}: {issue.message}")
        report.append("")
    
    if info:
        report.append("INFORMATION:")
        report.append("-" * 70)
        for issue in info:
            para_info = f" (Paragraph {issue.paragraph_number})" if issue.paragraph_number else ""
            report.append(f"  [{issue.section}]{para_info}: {issue.message}")
        report.append("")
    
    report.append("=" * 70)
    
    return "\n".join(report)


def create_test_document(filename: str, valid: bool = True):
    """Create a test .docx document for validation"""
    doc = Document()
    
    if valid:
        # Create valid memorandum
        doc.add_paragraph("12 March 2025")
        doc.add_paragraph()
        doc.add_paragraph("MEMORANDUM FOR RECORD")
        doc.add_paragraph()
        doc.add_paragraph("FROM: AFROTC/CC")
        doc.add_paragraph()
        doc.add_paragraph("SUBJECT: Put Subject Here")
        doc.add_paragraph()
        doc.add_paragraph("1. This memorandum provides ...")
        doc.add_paragraph()
        doc.add_paragraph("2. For Questions regarding Memo -- please use the T&Q.")
        doc.add_paragraph()
        doc.add_paragraph("3. Additional information can be found in the attachments.")
        doc.add_paragraph()
        doc.add_paragraph("//SIGNED//")
        doc.add_paragraph()
        doc.add_paragraph("Snuff A. Brown, Colonel, USAF")
        doc.add_paragraph()
        doc.add_paragraph("Commander")
        doc.add_paragraph()
        doc.add_paragraph("Attachments:")
        doc.add_paragraph()
        doc.add_paragraph("Tab 1")
        doc.add_paragraph()
        doc.add_paragraph("Tab 2")
    else:
        # Create invalid memorandum with issues
        doc.add_paragraph("March 12, 2025")  # Wrong date format
        doc.add_paragraph()
        doc.add_paragraph("MEMO FOR RECORD")  # Wrong header
        doc.add_paragraph()
        doc.add_paragraph("FROM AFROTC/CC")  # Missing colon
        doc.add_paragraph()
        doc.add_paragraph("SUBJECT: Testing")
        doc.add_paragraph()
        doc.add_paragraph("This is paragraph one but not numbered.")  # Not numbered
        doc.add_paragraph()
        doc.add_paragraph("This is paragraph two but not numbered.")  # Not numbered
        doc.add_paragraph()
        doc.add_paragraph("//SIGNED//")
        doc.add_paragraph()
        doc.add_paragraph("Snuff Brown")  # Missing rank and branch
        doc.add_paragraph()
        doc.add_paragraph("Attachments:")
        # No tabs listed
    
    doc.save(filename)


# Example usage
if __name__ == "__main__":
    import sys
    
    # Create test documents
    print("Creating test documents...")
    create_test_document('test_memo_valid.docx', valid=True)
    create_test_document('test_memo_invalid.docx', valid=False)
    print("Test documents created: test_memo_valid.docx, test_memo_invalid.docx\n")
    
    # Validate valid document
    print("=" * 70)
    print("TESTING VALID MEMORANDUM")
    print("=" * 70)
    
    validator = MemorandumValidator()
    passed, issues = validator.validate_file('test_memo_valid.docx')
    print(generate_report(passed, issues))
    
    # Validate invalid document
    print("\n\n" + "=" * 70)
    print("TESTING INVALID MEMORANDUM")
    print("=" * 70)
    
    validator2 = MemorandumValidator()
    passed2, issues2 = validator2.validate_file('test_memo_invalid.docx')
    print(generate_report(passed2, issues2))
    
    # If user provides a file path as argument
    if len(sys.argv) > 1:
        print("\n\n" + "=" * 70)
        print(f"VALIDATING USER FILE: {sys.argv[1]}")
        print("=" * 70)
        
        validator3 = MemorandumValidator()
        passed3, issues3 = validator3.validate_file(sys.argv[1])
        print(generate_report(passed3, issues3))
